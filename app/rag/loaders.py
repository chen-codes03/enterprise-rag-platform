"""文档加载器。

支持 Markdown / PDF / Word 三种格式解析，返回 LangChain Document。
按扩展名分发，统一 metadata（source / format / page）。
"""
from __future__ import annotations

from pathlib import Path

import markdown
import docx
from bs4 import BeautifulSoup
from langchain_core.documents import Document
from pypdf import PdfReader


def load_markdown(path: str | Path) -> list[Document]:
    """解析 Markdown，去除语法符号得到纯文本。"""
    text = Path(path).read_text(encoding="utf-8")
    html = markdown.markdown(text)
    plain = BeautifulSoup(html, "html.parser").get_text(separator="\n").strip()
    return [
        Document(
            page_content=plain,
            metadata={"source": str(path), "format": "markdown"},
        )
    ]


def load_pdf(path: str | Path) -> list[Document]:
    """解析 PDF，按页返回 Document。"""
    reader = PdfReader(str(path))
    docs: list[Document] = []
    for i, page in enumerate(reader.pages):
        docs.append(
            Document(
                page_content=page.extract_text() or "",
                metadata={"source": str(path), "format": "pdf", "page": i + 1},
            )
        )
    return docs


def load_docx(path: str | Path) -> list[Document]:
    """解析 Word 文档，拼接段落文本。"""
    document = docx.Document(str(path))
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    return [
        Document(
            page_content=text,
            metadata={"source": str(path), "format": "docx"},
        )
    ]


_LOADERS = {
    ".md": load_markdown,
    ".markdown": load_markdown,
    ".pdf": load_pdf,
    ".docx": load_docx,
}


def load_document(path: str | Path) -> list[Document]:
    """按文件扩展名分发到对应加载器。"""
    ext = Path(path).suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"不支持的文档格式: {ext}，可选: {list(_LOADERS)}")
    return loader(path)
